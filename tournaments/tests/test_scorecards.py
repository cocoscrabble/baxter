from io import BytesIO

from django.test import SimpleTestCase, TestCase, tag
from django.urls import reverse
from docx import Document

from tournaments.scorecards import (
    ScorecardResult,
    ScorecardSpec,
    build_document,
    make_rounds,
    render_scorecards,
)
from tournaments.models import DivisionSettings, Pairing, ResultSlip
from tournaments.tests.test_views import setUpTournament

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _spec(player_name, n_rounds=6, **kwargs):
    return ScorecardSpec(
        tournament_name="Spring Open",
        tournament_date="May 28, 2026",
        player_name=player_name,
        rounds=make_rounds(range(1, n_rounds + 1)),
        **kwargs,
    )


class MakeRoundsTests(SimpleTestCase):
    def test_first_round_is_undivided_by_default(self):
        rounds = make_rounds(range(1, 5))
        self.assertFalse(rounds[0].divided)
        self.assertTrue(all(r.divided for r in rounds[1:]))

    def test_undivided_first_can_be_disabled(self):
        rounds = make_rounds(range(1, 5), undivided_first=False)
        self.assertTrue(all(r.divided for r in rounds))


class ScorecardStructureTests(SimpleTestCase):
    def test_each_round_is_two_grid_rows(self):
        doc = build_document([_spec("Alice", n_rounds=6)])
        table = doc.tables[0]
        # header row + two rows per round
        self.assertEqual(len(table.rows), 1 + 2 * 6)

    def test_round_number_column_is_always_merged(self):
        doc = build_document([_spec("Alice", n_rounds=6)])
        table = doc.tables[0]
        # The two grid rows of round 2 (divided) share one Round-number cell.
        self.assertIs(table.cell(3, 0)._tc, table.cell(4, 0)._tc)

    def test_undivided_round_merges_all_columns(self):
        doc = build_document([_spec("Alice", n_rounds=6)])
        table = doc.tables[0]
        # Round 1 (rows 1-2) is undivided: every data column is one merged cell.
        for col in range(1, 7):
            self.assertIs(table.cell(1, col)._tc, table.cell(2, col)._tc)

    def test_divided_round_splits_only_the_last_column(self):
        doc = build_document([_spec("Alice", n_rounds=6)])
        table = doc.tables[0]
        # Round 2 (rows 3-4): columns 0-5 are merged into one cell each, and
        # only the last column (Spread) keeps two separate rows.
        for col in range(0, 6):
            self.assertIs(table.cell(3, col)._tc, table.cell(4, col)._tc)
        self.assertIsNot(table.cell(3, 6)._tc, table.cell(4, 6)._tc)

    def test_player_name_appears(self):
        doc = build_document([_spec("Carol Danvers")])
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Carol Danvers", text)

    def test_rounds_split_across_pages(self):
        # 20 rounds exceeds the 14-round first page, so it splits into two tables.
        doc = build_document([_spec("Alice", n_rounds=20)])
        self.assertEqual(len(doc.tables), 2)

    def test_one_set_of_tables_per_player(self):
        doc = build_document([_spec("Alice", n_rounds=6), _spec("Bob", n_rounds=6)])
        self.assertEqual(len(doc.tables), 2)


class OpponentPrefillTests(SimpleTestCase):
    def test_supplied_opponents_fill_their_round_only(self):
        doc = build_document(
            [_spec("Alice", n_rounds=6, opponents={1: "Bob", 3: "Carol"})]
        )
        table = doc.tables[0]
        # Opponent column is index 1; round N's row pair starts at 1 + 2*(N-1).
        self.assertEqual(table.cell(1, 1).text, "Bob")     # round 1
        self.assertEqual(table.cell(5, 1).text, "Carol")   # round 3
        self.assertEqual(table.cell(3, 1).text, "")        # round 2 left blank

    def test_no_opponents_leaves_column_blank(self):
        doc = build_document([_spec("Alice", n_rounds=6)])
        table = doc.tables[0]
        for round_idx in range(6):
            self.assertEqual(table.cell(1 + 2 * round_idx, 1).text, "")

    def test_each_player_keeps_their_own_opponents(self):
        specs = [
            _spec("Alice", opponents={1: "Bob"}),
            _spec("Bob", opponents={1: "Alice"}),
        ]
        doc = build_document(specs)
        self.assertEqual(doc.tables[0].cell(1, 1).text, "Bob")
        self.assertEqual(doc.tables[1].cell(1, 1).text, "Alice")

    def test_prefilled_division_still_shares_images(self):
        # Opponents differ per player but the clone path must still hold, so the
        # QR + logo stay a single shared pair.
        specs = [
            _spec(f"P{i}", opponents={1: f"Opp{i}"}, qr_url="https://x.test/live")
            for i in range(4)
        ]
        doc = Document(BytesIO(render_scorecards(specs)))
        self.assertEqual(len(doc.part.package.image_parts._image_parts), 2)


_VML_OVAL = "{urn:schemas-microsoft-com:vml}oval"


def _circle_offsets(el):
    """Left margin (pt) of every VML ellipse-circle under ``el``, in order."""
    import re

    offsets = []
    for oval in el.iter(_VML_OVAL):
        m = re.search(r"margin-left:(-?[\d.]+)pt", oval.get("style", ""))
        offsets.append(float(m.group(1)))
    return offsets


def _first_pt():
    from tournaments.scorecards import CIRCLE_FIRST_H_OFFSET, _EMU_PER_PT
    return round(CIRCLE_FIRST_H_OFFSET / _EMU_PER_PT, 2)


def _second_pt():
    from tournaments.scorecards import CIRCLE_SECOND_H_OFFSET, _EMU_PER_PT
    return round(CIRCLE_SECOND_H_OFFSET / _EMU_PER_PT, 2)


class StartPrefillTests(SimpleTestCase):
    # Round N's row pair starts at row 1 + 2*(N-1); the Round cell is column 0.
    @staticmethod
    def _round_cell(table, round_number):
        return table.cell(1 + 2 * (round_number - 1), 0)

    def test_supplied_starts_circle_the_right_ordinal(self):
        doc = build_document(
            [_spec("Alice", n_rounds=6, starts={1: "1st", 3: "2nd"})]
        )
        table = doc.tables[0]
        # The prompt text is untouched; the seat is shown by a circle over it.
        self.assertIn("1st", self._round_cell(table, 1).text)
        self.assertIn("2nd", self._round_cell(table, 1).text)
        self.assertEqual(_circle_offsets(self._round_cell(table, 1)._tc),
                         [_first_pt()])
        self.assertEqual(_circle_offsets(self._round_cell(table, 3)._tc),
                         [_second_pt()])

    def test_unmarked_rounds_get_no_circle(self):
        doc = build_document([_spec("Alice", n_rounds=6, starts={1: "1st"})])
        table = doc.tables[0]
        self.assertEqual(_circle_offsets(self._round_cell(table, 2)._tc), [])
        # Exactly one circle in the whole card (round 1 only).
        self.assertEqual(len(_circle_offsets(table._tbl)), 1)

    def test_each_player_circles_their_own_seat(self):
        specs = [
            _spec("Alice", starts={1: "1st"}),
            _spec("Bob", starts={1: "2nd"}),
        ]
        doc = build_document(specs)
        self.assertEqual(_circle_offsets(self._round_cell(doc.tables[0], 1)._tc),
                         [_first_pt()])
        self.assertEqual(_circle_offsets(self._round_cell(doc.tables[1], 1)._tc),
                         [_second_pt()])

    def test_marked_division_still_shares_images(self):
        # Starts differ per player but the clone path must still hold.
        specs = [
            _spec(f"P{i}", starts={1: "1st" if i % 2 else "2nd"},
                  qr_url="https://x.test/live")
            for i in range(4)
        ]
        doc = Document(BytesIO(render_scorecards(specs)))
        self.assertEqual(len(doc.part.package.image_parts._image_parts), 2)

    def test_circles_are_wps_wrapped_with_a_vml_fallback(self):
        # A bare wps:wsp shape makes Word for the web flag the document corrupt;
        # a pure-VML ellipse it mis-positions. So each seat circle is a wps shape
        # inside mc:Choice (correct position, no corruption) paired with a VML
        # <v:oval> fallback, and no ellipse anchor may sit outside an mc:Choice.
        from docx.oxml.ns import qn

        mc = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        doc = build_document([_spec("Alice", starts={1: "1st", 3: "2nd"})])
        body = doc.element.body

        ellipse_choices = [
            c for c in body.iter(f"{{{mc}}}Choice")
            if any(g.get("prst") == "ellipse" for g in c.iter(qn("a:prstGeom")))
        ]
        wps_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        self.assertEqual(len(ellipse_choices), 2)
        for choice in ellipse_choices:
            self.assertEqual(choice.get("Requires"), "wps")
            # a:graphicData/@uri must name the wps:wsp child's namespace, or Word
            # (unlike OpenXmlValidator) can't resolve the shape and flags the doc
            # corrupt. Pinning it here guards against reintroducing that mismatch.
            gdata = choice.find(f".//{qn('a:graphicData')}")
            self.assertEqual(gdata.get("uri"), wps_ns)
            self.assertIsNotNone(gdata.find(f"{{{wps_ns}}}wsp"))
            fallback = choice.getparent().find(f"{{{mc}}}Fallback")
            self.assertIsNotNone(fallback)
            self.assertIsNotNone(fallback.find(f".//{_VML_OVAL}"))

        for anchor in body.iter(qn("wp:anchor")):
            if any(g.get("prst") == "ellipse" for g in anchor.iter(qn("a:prstGeom"))):
                ancestors = {a.tag for a in anchor.iterancestors()}
                self.assertIn(f"{{{mc}}}Choice", ancestors)


class ResultPrefillTests(SimpleTestCase):
    def test_record_and_scores_and_both_spread_subcells(self):
        # Round 2 is divided, so its Spread column has two subcells.
        doc = build_document([_spec("Alice", n_rounds=6, results={
            2: ScorecardResult(player_score=450, opponent_score=380,
                               cumulative_wins=2, cumulative_losses=1,
                               cumulative_spread=120),
        })])
        table = doc.tables[0]
        top, bottom = 3, 4  # round 2 grid rows (1 + 2*1, and the one below)
        self.assertEqual(table.cell(top, 2).text, "2")    # Won (running wins)
        self.assertEqual(table.cell(top, 3).text, "1")    # Lost (running losses)
        self.assertEqual(table.cell(top, 4).text, "450")  # Player Score
        self.assertEqual(table.cell(top, 5).text, "380")  # Opponent Score
        self.assertEqual(table.cell(top, 6).text, "70")   # game spread
        self.assertEqual(table.cell(bottom, 6).text, "120")  # cumulative spread

    def test_half_records_from_a_tie_show_a_fraction(self):
        doc = build_document([_spec("Alice", n_rounds=6, results={
            2: ScorecardResult(player_score=400, opponent_score=400,
                               cumulative_wins=2.5, cumulative_losses=1.5,
                               cumulative_spread=0),
        })])
        table = doc.tables[0]
        self.assertEqual(table.cell(3, 2).text, "2.5")  # Won
        self.assertEqual(table.cell(3, 3).text, "1.5")  # Lost

    def test_rounds_without_a_result_stay_blank(self):
        doc = build_document([_spec("Alice", n_rounds=6, results={
            2: ScorecardResult(player_score=450, opponent_score=380,
                               cumulative_wins=1, cumulative_spread=70),
        })])
        table = doc.tables[0]
        # Round 3 (rows 5/6) has no result: score columns are empty.
        for col in (2, 3, 4, 5, 6):
            self.assertEqual(table.cell(5, col).text, "")


class RenderScorecardsTests(SimpleTestCase):
    def test_returns_openable_docx_bytes(self):
        data = render_scorecards([_spec("Alice")])
        self.assertTrue(data.startswith(b"PK"))  # docx is a zip archive
        Document(BytesIO(data))  # parses without error

    def test_qr_url_embeds_an_image(self):
        data = render_scorecards([_spec("Alice", qr_url="https://example.test/live")])
        doc = Document(BytesIO(data))
        # logo + QR = two images; without a qr_url only the logo is present.
        self.assertEqual(len(doc.inline_shapes) + _floating_image_count(doc), 2)


def _floating_image_count(doc):
    from docx.oxml.ns import qn

    return len(doc.element.findall(".//" + qn("a:blip")))


class ClonedScorecardTests(SimpleTestCase):
    """The division path clones one template per player; guard its correctness."""

    def test_every_player_gets_their_own_name(self):
        names = ["Alice Walker", "Bob Stevens", "Carol Ng"]
        specs = [_spec(n) for n in names]
        doc = Document(BytesIO(render_scorecards(specs)))
        text = "\n".join(p.text for p in doc.paragraphs)
        for name in names:
            self.assertIn(name, text)

    def test_shared_images_are_embedded_once(self):
        # QR + logo are identical across the division, so cloning should reuse
        # the same two image parts no matter how many players there are.
        specs = [_spec(f"Player {i}", qr_url="https://x.test/live") for i in range(5)]
        doc = Document(BytesIO(render_scorecards(specs)))
        self.assertEqual(len(doc.part.package.image_parts._image_parts), 2)

    def test_cloned_drawings_have_unique_ids(self):
        from docx.oxml.ns import qn

        specs = [_spec(f"Player {i}", qr_url="https://x.test/live") for i in range(4)]
        doc = Document(BytesIO(render_scorecards(specs)))
        # Both the wordprocessing docPr id and the DrawingML picture id must be
        # unique across the package; a duplicate in either makes Word for the web
        # flag the file corrupt and drop the images to placeholders (add_picture
        # emits every pic:cNvPr id as 0, so the clones must renumber them).
        for tag in ("wp:docPr", "pic:cNvPr"):
            ids = [e.get("id") for e in doc.element.iter(qn(tag))]
            self.assertEqual(len(ids), len(set(ids)), f"{tag} ids not unique: {ids}")

    def test_differing_layouts_fall_back_and_still_render(self):
        specs = [
            _spec("Alice", n_rounds=6),
            ScorecardSpec("Other Cup", "May 28, 2026", "Bob",
                          rounds=make_rounds(range(1, 9))),
        ]
        doc = Document(BytesIO(render_scorecards(specs)))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Alice", text)
        self.assertIn("Bob", text)
        self.assertIn("Other Cup", text)


@tag("slow")
class DivisionScorecardsViewTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        setUpTournament(cls)

    def setUp(self):
        # Scorecards are an organiser tool — editor-only.
        self.client.login(username="owner", password="testpass123")

    def test_downloads_docx_attachment(self):
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], DOCX_CONTENT_TYPE)
        self.assertIn("attachment", response["Content-Disposition"])
        self.assertIn(".docx", response["Content-Disposition"])
        self.assertTrue(response.content.startswith(b"PK"))

    def test_filename_is_slugified(self):
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
        )
        self.assertIn("test-tournament-open-scorecards.docx",
                      response["Content-Disposition"])

    def test_forbidden_for_non_editor(self):
        # The whole scorecards feature is editor-only.
        self.client.login(username="other", password="testpass123")
        for name in ("division_scorecards", "division_scorecards_download"):
            response = self.client.get(
                reverse(name, kwargs=self.division.slug_kwargs())
            )
            self.assertEqual(response.status_code, 403)

    def test_opponents_prefilled_from_pairings(self):
        # Keep it to 3 rounds so each player's card is a single table.
        DivisionSettings.objects.create(
            division=self.division,
            round_pairings=[{"round": 1}, {"round": 2}, {"round": 3}],
        )
        Pairing.objects.create(
            division=self.division, round=1,
            first=self.entrant1, second=self.entrant2,
        )
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
        )
        doc = Document(BytesIO(response.content))
        # Entrants order by number: table[0] is player1's card, table[1] player2's.
        # Round 1's Opponent cell is row 1, column 1.
        self.assertEqual(doc.tables[0].cell(1, 1).text, self.player2.name)
        self.assertEqual(doc.tables[1].cell(1, 1).text, self.player1.name)

    def _three_round_result(self):
        DivisionSettings.objects.create(
            division=self.division,
            round_pairings=[{"round": 1}, {"round": 2}, {"round": 3}],
        )
        ResultSlip.objects.create(
            division=self.division, round=1,
            winner=self.entrant1, winner_score=450,
            loser=self.entrant2, loser_score=380,
            winner_started=True,
        )

    def test_results_prefilled_when_requested(self):
        self._three_round_result()
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
            + "?include_results=1"
        )
        doc = Document(BytesIO(response.content))
        # Round 1 (undivided) is a single merged block; its cells are row 1.
        winner_card, loser_card = doc.tables[0], doc.tables[1]
        # After round 1 the winner's running record is 1-0, the loser's 0-1.
        self.assertEqual(winner_card.cell(1, 2).text, "1")    # Won
        self.assertEqual(winner_card.cell(1, 3).text, "0")    # Lost
        self.assertEqual(winner_card.cell(1, 4).text, "450")  # Player Score
        self.assertEqual(winner_card.cell(1, 5).text, "380")  # Opponent Score
        self.assertEqual(loser_card.cell(1, 2).text, "0")     # Won
        self.assertEqual(loser_card.cell(1, 3).text, "1")     # Lost
        self.assertEqual(loser_card.cell(1, 4).text, "380")

    def test_results_omitted_without_the_flag(self):
        self._three_round_result()
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
        )
        doc = Document(BytesIO(response.content))
        # Score columns stay blank when results aren't requested.
        self.assertEqual(doc.tables[0].cell(1, 2).text, "")
        self.assertEqual(doc.tables[0].cell(1, 4).text, "")

    def test_starts_circled_from_pairings(self):
        DivisionSettings.objects.create(
            division=self.division,
            round_pairings=[{"round": 1}, {"round": 2}, {"round": 3}],
        )
        Pairing.objects.create(
            division=self.division, round=1,
            first=self.entrant1, second=self.entrant2,
        )
        response = self.client.get(
            reverse("division_scorecards_download", kwargs=self.division.slug_kwargs())
        )
        doc = Document(BytesIO(response.content))
        # entrant1 went first, entrant2 second; round 1's Round cell is (1, 0).
        self.assertEqual(_circle_offsets(doc.tables[0].cell(1, 0)._tc),
                         [_first_pt()])
        self.assertEqual(_circle_offsets(doc.tables[1].cell(1, 0)._tc),
                         [_second_pt()])


# Canonical child orders for the OOXML complex types we emit by hand. Word
# validates each against these sequences on open and rejects the file if a
# child appears out of order; LibreOffice silently repairs it, so only a schema
# check (not "does it open here") catches the regression. Local names only.
_CT_TBL_PR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
    "tblDescription", "tblPrChange",
]
_CT_TC_PR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
    "cellIns", "cellDel", "cellMerge", "tcPrChange",
]
_CT_ANCHOR_ORDER = [
    "simplePos", "positionH", "positionV", "extent", "effectExtent",
    # any one wrap* variant occupies this slot
    "wrapNone", "wrapSquare", "wrapTight", "wrapThrough", "wrapTopAndBottom",
    "docPr", "cNvGraphicFramePr", "graphic", "sizeRelH", "sizeRelV",
]
# the wrap variants are mutually exclusive alternatives sharing one rank
_WRAP_RANK = {w: 5 for w in _CT_ANCHOR_ORDER[5:10]}


def _local(tag):
    return tag.rsplit("}", 1)[-1]


class SchemaOrderTests(SimpleTestCase):
    """Every hand-built tblPr / tcPr / anchor must be in OOXML sequence order.

    LibreOffice tolerates out-of-order children; Word raises "Word experienced
    an error trying to open the file". These walk the generated document and
    fail if any child is out of its type's canonical order.
    """

    @staticmethod
    def _rich_doc():
        # Exercise every raw-XML path: borders + fixed layout (tblPr), shading +
        # width + merges + vAlign (tcPr), and floating QR/logo/ellipse (anchor).
        return build_document([
            _spec("Alice", n_rounds=6, opponents={1: "Bob"},
                  starts={1: "1st", 3: "2nd"}, qr_url="https://x.test/live"),
        ])

    def _assert_ordered(self, el, order, *, strict):
        rank = {name: i for i, name in enumerate(order)}
        rank.update(_WRAP_RANK if order is _CT_ANCHOR_ORDER else {})
        last = -1
        for child in el:
            name = _local(child.tag)
            if name not in rank:
                if strict:
                    self.fail(
                        f"<{name}> not in canonical order for this element; "
                        "the test's schema list needs updating"
                    )
                continue
            self.assertGreaterEqual(
                rank[name], last,
                f"<{name}> is out of OOXML sequence order (Word will reject it)",
            )
            last = rank[name]

    def test_tblPr_children_in_order(self):
        from docx.oxml.ns import qn
        body = self._rich_doc().element.body
        tblPrs = list(body.iter(qn("w:tblPr")))
        self.assertTrue(tblPrs)
        for tblPr in tblPrs:
            self._assert_ordered(tblPr, _CT_TBL_PR_ORDER, strict=True)

    def test_tcPr_children_in_order(self):
        from docx.oxml.ns import qn
        body = self._rich_doc().element.body
        tcPrs = list(body.iter(qn("w:tcPr")))
        self.assertTrue(tcPrs)
        for tcPr in tcPrs:
            self._assert_ordered(tcPr, _CT_TC_PR_ORDER, strict=True)

    def test_anchor_children_in_order(self):
        from docx.oxml.ns import qn
        body = self._rich_doc().element.body
        anchors = list(body.iter(qn("wp:anchor")))
        self.assertTrue(anchors)
        for anchor in anchors:
            self._assert_ordered(anchor, _CT_ANCHOR_ORDER, strict=False)
