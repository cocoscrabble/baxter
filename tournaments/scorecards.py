"""Generate printable player scorecards as a Word (.docx) document.

This module is deliberately free of Django imports so it can be exercised in
isolation (tests, management commands, future async jobs). Callers build a list
of :class:`ScorecardSpec` — one per player — and hand it to
:func:`render_scorecards`, which returns the ``.docx`` file as bytes.

Layout per player: a centred title block (tournament name / date / player name)
with the COCO logo floated left and a QR code (encoding the live-coverage URL)
floated right, followed by one or more bordered round tables and a footer. Each
round is two grid rows so the player can record two games; the Round-number
cell is vertically merged so the label shows once. A round marked ``divided=
False`` (round 1) merges every column instead, leaving a single open block with
no internal divider.
"""

from __future__ import annotations

import copy
import os
from dataclasses import dataclass, field, replace
from io import BytesIO

import qrcode
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.table import _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsmap, qn
from docx.shared import Emu, Pt

FONT = "Source Sans Pro"
HEADER_FILL = "f2f2f2"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "coco_logo.png")

# Column widths in twentieths of a point (dxa).
COL_WIDTHS = [1081, 2604, 900, 900, 1620, 1710, 1530]
HEADERS = ["Round", "Opponent", "Won", "Lost",
           "Player Score", "Opponent Score", "Spread"]
ROW_HEIGHT = 346  # dxa; every grid row uses this so the table stays uniform.

# Image geometry (EMU), reproduced from the reference scorecard.
LOGO_SIZE = 840658
QR_SIZE = 840105
LOGO_H_OFFSET, LOGO_V_OFFSET = 1, -634
QR_H_OFFSET, QR_V_OFFSET = 5789295, 7315

# Ellipse drawn around the "1st"/"2nd" prompt to mark which seat the player took.
# Offsets (EMU) are relative to the Round cell's column / the prompt paragraph;
# the "1st" sits left of centre and "2nd" right, in a ~54pt-wide cell. Tune these
# if a font change shifts where the ordinals land.
CIRCLE_COLOR = "531882"  # outline colour (purple), like a pen mark
CIRCLE_W, CIRCLE_H = 295000, 200000
# Centred over the ordinal ink: the "1st"/"2nd" text sits at roughly ∓40pt from
# the cell centre, so the ellipse's left edge (posOffset) is set to land its
# centre there. The offset is paragraph-relative (from the prompt paragraph's
# top), tuned to centre the ellipse in desktop Word and LibreOffice. Word for the
# web renders the same offset a touch high, but that's a Word-online rendering
# glitch — the printed/desktop output is what matters, so we optimise for it.
CIRCLE_V_OFFSET = -30000
CIRCLE_FIRST_H_OFFSET = -25000  # nudged slightly left to sit over "1st"
CIRCLE_SECOND_H_OFFSET = 245000


@dataclass(frozen=True)
class RoundSpec:
    """One round on the scorecard.

    ``divided`` rounds get two writable lines (a horizontal divider between
    them); undivided rounds present a single open block of the same height.
    """

    number: int
    divided: bool = True


@dataclass(frozen=True)
class ScorecardResult:
    """A submitted game result, from one player's point of view, used to
    prefill the score columns for a round."""

    player_score: int
    opponent_score: int
    # Running record through this round, for the Won / Lost columns. A win
    # counts 1, a loss 0, a tie half to each side (so these can be half-integers).
    cumulative_wins: float = 0
    cumulative_losses: float = 0
    # Running spread total through this round, for the lower Spread subcell.
    cumulative_spread: int = 0

    @property
    def spread(self) -> int:
        """This round's spread (goes in the upper Spread subcell)."""
        return self.player_score - self.opponent_score


@dataclass(frozen=True)
class ScorecardSpec:
    """Everything needed to render one player's scorecard."""

    tournament_name: str
    tournament_date: str
    player_name: str
    rounds: list[RoundSpec]
    # Optional {round number: opponent name} to prefill the Opponent column.
    opponents: dict[int, str] = field(default_factory=dict)
    # Optional {round number: "1st"/"2nd"} marking whether the player went
    # first or second that round, where the pairing fixes it.
    starts: dict[int, str] = field(default_factory=dict)
    # Optional {round number: ScorecardResult} prefilling the score columns for
    # rounds whose result has already been submitted.
    results: dict[int, "ScorecardResult"] = field(default_factory=dict)
    qr_url: str = ""
    footer_text: str = "Submit results and view pairings and standings at:"
    footer_url: str = "cocoscrabble.org/live-coverage"
    # Rounds before the first page break, and per continuation page thereafter.
    first_page_rounds: int = 14
    rounds_per_page: int = 20


def make_rounds(round_numbers, *, undivided_first=True):
    """Build a RoundSpec list from round numbers, optionally leaving round 1 open."""
    specs = []
    for i, n in enumerate(round_numbers):
        specs.append(RoundSpec(number=n, divided=not (undivided_first and i == 0)))
    return specs


# --- low-level docx helpers -------------------------------------------------

def _set_run_font(run, size, *, bold=False):
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def _style_paragraph(para):
    pf = para.paragraph_format
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _add_paragraph(doc, text, size, *, bold=False, center=True):
    p = doc.add_paragraph()
    _style_paragraph(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        _set_run_font(p.add_run(text), size, bold=bold)
    return p


def _set_cell_shading(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    # w:shd must precede everything after it in the CT_TcPr sequence, or Word
    # rejects the file (LibreOffice tolerates out-of-order children).
    cell._tc.get_or_add_tcPr().insert_element_before(
        shd,
        "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
        "w:vAlign", "w:hideMark", "w:cellIns", "w:cellDel",
        "w:cellMerge", "w:tcPrChange",
    )


def _set_cell_width(cell, dxa):
    # get_or_add_tcW places w:tcW at its correct CT_TcPr sequence position.
    tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")


def _set_grid_widths(table, widths):
    for col, width in zip(table._tbl.tblGrid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(width))


def _set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    # w:tblBorders must precede w:tblLayout (and the rest) in the CT_TblPr
    # sequence; appending blindly puts them out of order and Word refuses to
    # open the file, while LibreOffice repairs it silently.
    table._tbl.tblPr.insert_element_before(
        borders,
        "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
        "w:tblCaption", "w:tblDescription", "w:tblPrChange",
    )


def _set_fixed_layout(table):
    # get_or_add_tblLayout places w:tblLayout at its correct sequence position.
    table._tbl.tblPr.get_or_add_tblLayout().set(qn("w:type"), "fixed")


def _set_row_height(row, dxa):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(dxa))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


def _clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def _center_cell(cell):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _style_paragraph(cell.paragraphs[0])
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fill_header_cell(cell, text):
    _set_cell_shading(cell, HEADER_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    _style_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(text), 10, bold=True)


def _fill_round_cell(cell, round_number, *, mark_start=False):
    """Round number on the first line, '1st   2nd' (superscript) below it.

    The '1st'/'2nd' prompt is always written out. When ``mark_start`` is set the
    round's seat may be fixed for some player, so a placeholder run is appended;
    :func:`_resolve_start` later turns it into an ellipse over the right ordinal
    (or removes it) per player.
    """
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _clear_cell(cell)

    p1 = cell.add_paragraph()
    _style_paragraph(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p1.add_run(str(round_number)), 10)

    p2 = cell.add_paragraph()
    _style_paragraph(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, sup in [("1", False), ("st", True), ("   2", False),
                      ("nd", True), (" ", False)]:
        run = p2.add_run(text)
        _set_run_font(run, 10)
        if sup:
            run.font.superscript = True
    if mark_start:
        _set_run_font(p2.add_run(_start_sentinel(round_number)), 1)


_uid = [1000]


def _add_floating_image(paragraph, image, cx, cy, h_offset, v_offset, *, behind):
    """Anchor a floating image (path or stream) at an absolute offset in EMU."""
    run = paragraph.add_run()
    run.add_picture(image, width=Emu(cx), height=Emu(cy))
    inline = run._r.find(qn("w:drawing"))[0]
    graphic = inline.find(qn("a:graphic"))

    _uid[0] += 1
    anchor = parse_xml(
        f'<wp:anchor xmlns:wp="{nsmap["wp"]}" xmlns:a="{nsmap["a"]}" '
        f'xmlns:r="{nsmap["r"]}" xmlns:pic="{nsmap["pic"]}" '
        f'simplePos="0" relativeHeight="0" behindDoc="{1 if behind else 0}" '
        'locked="0" layoutInCell="1" allowOverlap="1" '
        'distT="0" distB="0" distL="0" distR="0">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{h_offset}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{v_offset}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{_uid[0]}" name="image{_uid[0]}"/>'
        '<wp:cNvGraphicFramePr/>'
        '</wp:anchor>'
    )
    anchor.append(graphic)
    # The DrawingML non-visual picture id (pic:cNvPr/@id) that add_picture emits
    # is always 0; leaving every logo/QR pic at 0 gives the package duplicate
    # drawing-object ids. Desktop Word and LibreOffice silently renumber, but
    # Word for the web rejects the file as corrupt and drops the images to
    # "unable to load picture" placeholders. Give each pic its own id.
    for cnvpr in anchor.iter(qn("pic:cNvPr")):
        cnvpr.set("id", str(_uid[0]))
    drawing = run._r.find(qn("w:drawing"))
    drawing.remove(inline)
    drawing.append(anchor)


def _make_qr_png(data):
    """Render ``data`` as a QR code and return PNG bytes."""
    img = qrcode.make(data)  # PilImage; its default kind is PNG
    buf = BytesIO()
    img.save(buf)
    return buf.getvalue()


# --- assembly ---------------------------------------------------------------

def _row_cells(row, table):
    """The row's cells straight off its ``<w:tc>`` children.

    ``_Row.cells`` / ``Table.cell`` rebuild the whole table grid on every
    access (O(rows*cols) each), which makes filling a large table quadratic.
    Reading ``tr.tc_lst`` is O(1) and there are no horizontal merges here, so
    each row's ``<w:tc>`` list is exactly its column cells.
    """
    return [_Cell(tc, table) for tc in row._tr.tc_lst]


def _set_vmerge(cell, *, restart):
    """Vertically merge this cell with the one below (restart) or into it (continue)."""
    vMerge = cell._tc.get_or_add_tcPr().get_or_add_vMerge()
    vMerge.val = "restart" if restart else None


# The seat circle is a DrawingML wps:wsp shape (Word 2010 extension) — Word for
# the web positions this correctly, unlike a VML shape — wrapped in
# mc:AlternateContent with a VML fallback, which is the only way Word writes such
# a shape and what keeps Word for the web from flagging the document corrupt.
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_WPS_DRAWING_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
# Namespaces for the VML fallback; python-docx's nsmap doesn't carry these, but
# the document root declares them so the parsed fragments resolve.
_V_NS = "urn:schemas-microsoft-com:vml"
_O_NS = "urn:schemas-microsoft-com:office:office"
_W10_NS = "urn:schemas-microsoft-com:office:word"

_EMU_PER_PT = 12700  # VML positions in points; our shape geometry is in EMU.

_QN_DOCPR = qn("wp:docPr")
_QN_ANCHOR = qn("wp:anchor")
_QN_ANCHOR_ID = f"{{{_WPS_DRAWING_NS}}}anchorId"
_QN_EDIT_ID = f"{{{_WPS_DRAWING_NS}}}editId"
_QN_VML_OVAL = f"{{{_V_NS}}}oval"
_QN_O_SPID = f"{{{_O_NS}}}spid"
# z-order base for the seat ellipses; each clone gets base+uid so no two shapes
# share a relativeHeight (Word assigns each floating shape a distinct one).
_CIRCLE_Z_BASE = 251659264

# A round whose seat may be fixed gets a placeholder run in its prompt paragraph
# carrying this prefix plus the round number. The per-player patch pass swaps
# each placeholder for an ellipse over the right ordinal — or removes it —
# riding the same walk that fills in the name and opponents, so locating the
# marks costs nothing extra.
_START_PREFIX = ""  # U+E000 (private use): never appears in real text


def _start_sentinel(round_number):
    return f"{_START_PREFIX}{round_number}"


def _circle_xml(h_offset):
    """A run holding the seat ellipse, wrapped Word's way for compatibility.

    The ellipse is a ``wps:wsp`` DrawingML shape (Word 2010 extension). Word for
    the web positions it correctly — but a *bare* ``wps`` shape (one Word never
    writes) makes it flag the whole document corrupt, and a pure-VML ellipse it
    renders in the wrong place (it ignores VML floating offsets). So we do what
    Word does: the ``wps`` shape is the ``mc:Choice`` (Word and Word-online read
    it — correct position, no corruption flag), with a VML ``<w:pict>`` oval as
    the ``mc:Fallback`` for clients that don't grok ``wps``.

    Word for the web is stricter than the OOXML schema: a *skeletal* ``wps``
    shape validates but still trips its corruption check. So the ``mc:Choice``
    mirrors, element for element, what desktop Word writes for a hand-drawn oval
    — populated ``wps:bodyPr``, ``wps:style``, ``a:effectLst``, and the
    ``wp14:sizeRel*`` extensions — differing only in our geometry, colour, and
    the paragraph-relative vertical anchor (each card is cloned down the page, so
    the ellipse must ride its own prompt paragraph rather than a page offset).
    The single ``<w:r>`` root declares every prefix the two branches use so the
    fragment resolves standalone.
    """
    drawing = (
        '<w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" '
        f'simplePos="0" relativeHeight="{_CIRCLE_Z_BASE}" behindDoc="0" '
        'locked="0" layoutInCell="1" allowOverlap="1" '
        'wp14:anchorId="00000000" wp14:editId="00000000">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{h_offset}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{CIRCLE_V_OFFSET}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{CIRCLE_W}" cy="{CIRCLE_H}"/>'
        '<wp:effectExtent l="0" t="0" r="9525" b="9525"/><wp:wrapNone/>'
        '<wp:docPr id="0" name="circle0"/><wp:cNvGraphicFramePr/>'
        '<a:graphic><a:graphicData uri="' + _WPS_NS + '">'
        '<wps:wsp><wps:cNvSpPr/><wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{CIRCLE_W}" cy="{CIRCLE_H}"/></a:xfrm>'
        '<a:prstGeom prst="ellipse"><a:avLst/></a:prstGeom><a:noFill/>'
        f'<a:ln w="19050"><a:solidFill><a:srgbClr val="{CIRCLE_COLOR}"/></a:solidFill></a:ln>'
        '<a:effectLst/></wps:spPr>'
        '<wps:style>'
        '<a:lnRef idx="1"><a:schemeClr val="accent1"/></a:lnRef>'
        '<a:fillRef idx="3"><a:schemeClr val="accent1"/></a:fillRef>'
        '<a:effectRef idx="2"><a:schemeClr val="accent1"/></a:effectRef>'
        '<a:fontRef idx="minor"><a:schemeClr val="lt1"/></a:fontRef>'
        '</wps:style>'
        '<wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow" '
        'horzOverflow="overflow" vert="horz" wrap="square" lIns="91440" '
        'tIns="45720" rIns="91440" bIns="45720" numCol="1" spcCol="0" '
        'rtlCol="0" fromWordArt="0" anchor="ctr" anchorCtr="0" forceAA="0" '
        'compatLnSpc="1"><a:prstTxWarp prst="textNoShape"><a:avLst/>'
        '</a:prstTxWarp><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp></a:graphicData></a:graphic>'
        '<wp14:sizeRelH relativeFrom="margin"><wp14:pctWidth>0</wp14:pctWidth></wp14:sizeRelH>'
        '<wp14:sizeRelV relativeFrom="margin"><wp14:pctHeight>0</wp14:pctHeight></wp14:sizeRelV>'
        '</wp:anchor></w:drawing>'
    )
    # VML fallback: same ellipse in points, positioned like the wps anchor's
    # column/paragraph origins. Only clients that can't read wps use this; its
    # style string mirrors what Word emits for the fallback branch.
    fallback = (
        '<w:pict>'
        '<v:oval id="circle0" o:spid="_x0000_s1026" '
        f'style="position:absolute;margin-left:{h_offset / _EMU_PER_PT:.2f}pt;'
        f'margin-top:{CIRCLE_V_OFFSET / _EMU_PER_PT:.2f}pt;'
        f'width:{CIRCLE_W / _EMU_PER_PT:.2f}pt;height:{CIRCLE_H / _EMU_PER_PT:.2f}pt;'
        f'z-index:{_CIRCLE_Z_BASE};visibility:visible;mso-wrap-style:square;'
        'mso-wrap-distance-left:9pt;mso-wrap-distance-top:0;'
        'mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:0;'
        'mso-position-horizontal:absolute;mso-position-horizontal-relative:text;'
        'mso-position-vertical:absolute;mso-position-vertical-relative:text;'
        'v-text-anchor:middle" '
        f'filled="f" strokecolor="#{CIRCLE_COLOR}" strokeweight="1.5pt">'
        '<w10:wrap type="none"/></v:oval></w:pict>'
    )
    return (
        f'<w:r xmlns:w="{nsmap["w"]}" xmlns:wp="{nsmap["wp"]}" '
        f'xmlns:a="{nsmap["a"]}" xmlns:wps="{_WPS_NS}" xmlns:mc="{_MC_NS}" '
        f'xmlns:wp14="{_WPS_DRAWING_NS}" '
        f'xmlns:v="{_V_NS}" xmlns:o="{_O_NS}" xmlns:w10="{_W10_NS}">'
        '<mc:AlternateContent>'
        f'<mc:Choice Requires="wps">{drawing}</mc:Choice>'
        f'<mc:Fallback>{fallback}</mc:Fallback>'
        '</mc:AlternateContent></w:r>'
    )


# One parsed ellipse run per side, cloned per insertion. parse_xml is ~3x the
# cost of a deepcopy and was being paid once per ellipse; building each side's
# XML once and copying instead keeps a big division's cards cheap.
_CIRCLE_RUNS = {}


def _circle_run(*, first):
    base = _CIRCLE_RUNS.get(first)
    if base is None:
        h = CIRCLE_FIRST_H_OFFSET if first else CIRCLE_SECOND_H_OFFSET
        base = parse_xml(_circle_xml(h))
        _CIRCLE_RUNS[first] = base
    run = copy.deepcopy(base)
    _uid[0] += 1
    uid = _uid[0]
    # The wps drawing id, its anchor/edit ids and z-order, and the VML fallback
    # shape id all have to be unique across the document, or Word rejects it.
    docPr = run.find(".//" + _QN_DOCPR)
    docPr.set("id", str(uid))
    docPr.set("name", f"circle{uid}")
    anchor = run.find(".//" + _QN_ANCHOR)
    anchor.set(_QN_ANCHOR_ID, f"{uid:08X}")
    anchor.set(_QN_EDIT_ID, f"{uid ^ 0x55555555:08X}")
    anchor.set("relativeHeight", str(_CIRCLE_Z_BASE + uid))
    oval = run.find(".//" + _QN_VML_OVAL)
    oval.set("id", f"circle{uid}")
    oval.set(_QN_O_SPID, f"_x0000_s{1026 + uid}")
    return run


def _resolve_start(run, round_number, spec):
    """Turn a start placeholder ``<w:r>`` into an ellipse, or drop it.

    ``spec.starts`` says whether the player took the 1st or 2nd seat that round;
    we replace the placeholder run with an ellipse over that ordinal, or remove
    it when the round has no fixed seat for this player.
    """
    para = run.getparent()
    seat = spec.starts.get(round_number)
    if seat:
        para.replace(run, _circle_run(first=seat == "1st"))
    else:
        para.remove(run)


def _prefill_cell(cell, text):
    """Write a 10pt run into a cell's first paragraph (no-op for empty text)."""
    if text != "":
        _set_run_font(cell.paragraphs[0].add_run(text), 10)


def _fmt_record(value):
    """Format a running win/loss count, dropping a trailing ``.0`` (so 3, not
    3.0) but keeping a half (3.5)."""
    return f"{value:.1f}".rstrip("0").rstrip(".")


def _fill_result(top, bottom, spec, result):
    """Prefill a round's score columns from a submitted result.

    The Won and Lost columns carry the player's running win/loss record through
    this round; the player's and opponent's scores go in their columns; the game
    spread goes in the upper Spread subcell and the cumulative spread in the
    lower one (for a divided round — an undivided round has a single, merged
    Spread cell, so only the game spread is shown there).
    """
    won, lost, pscore, oscore, spread = 2, 3, 4, 5, len(HEADERS) - 1
    _prefill_cell(top[won], _fmt_record(result.cumulative_wins))
    _prefill_cell(top[lost], _fmt_record(result.cumulative_losses))
    _prefill_cell(top[pscore], str(result.player_score))
    _prefill_cell(top[oscore], str(result.opponent_score))
    _prefill_cell(top[spread], str(result.spread))
    if spec.divided:
        _prefill_cell(bottom[spread], str(result.cumulative_spread))


def _add_round_table(doc, round_specs, opponents, results, placeholder_rounds):
    """One bordered table covering ``round_specs`` (each round = two grid rows).

    ``opponents`` is a {round number: name} mapping used to prefill the
    Opponent column; rounds absent from it are left blank. ``results`` is a
    {round number: ScorecardResult} mapping prefilling the score columns for
    rounds whose result is in. ``placeholder_rounds`` are the rounds whose
    Round cell gets a seat placeholder for later resolving.
    """
    table = doc.add_table(rows=1 + 2 * len(round_specs), cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_grid_widths(table, COL_WIDTHS)
    _set_table_borders(table)
    _set_fixed_layout(table)

    rows = table.rows
    _set_row_height(rows[0], ROW_HEIGHT)
    for cell, text, width in zip(_row_cells(rows[0], table), HEADERS, COL_WIDTHS):
        _set_cell_width(cell, width)
        _fill_header_cell(cell, text)

    r = 1
    for spec in round_specs:
        top = _row_cells(rows[r], table)
        bottom = _row_cells(rows[r + 1], table)
        _set_row_height(rows[r], ROW_HEIGHT)
        _set_row_height(rows[r + 1], ROW_HEIGHT)
        for width, top_cell, bottom_cell in zip(COL_WIDTHS, top, bottom):
            _set_cell_width(top_cell, width)
            _set_cell_width(bottom_cell, width)

        # Every column is vertically merged across the round's two rows except
        # the last (Spread), which stays split into two rows for divided rounds
        # (round 1 is undivided, so even Spread merges into one open block).
        last = len(HEADERS) - 1
        for c in range(len(HEADERS)):
            split = c == last and spec.divided
            if not split:
                _set_vmerge(top[c], restart=True)
                _set_vmerge(bottom[c], restart=False)
            _center_cell(top[c])
            _center_cell(bottom[c])

        # Round number goes in the (merged) first cell so it appears once.
        _fill_round_cell(top[0], spec.number,
                         mark_start=spec.number in placeholder_rounds)

        # Prefill the (merged) Opponent cell if a name was supplied.
        _prefill_cell(top[1], opponents.get(spec.number, ""))

        # Prefill the score columns if this round's result is in.
        result = results.get(spec.number)
        if result is not None:
            _fill_result(top, bottom, spec, result)
        r += 2

    return table


def _add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _chunk_rounds(spec):
    """Split rounds into per-page groups (smaller first page leaves room for the title)."""
    rounds = spec.rounds
    chunks = [rounds[: spec.first_page_rounds]]
    rest = rounds[spec.first_page_rounds:]
    for i in range(0, len(rest), spec.rounds_per_page):
        chunks.append(rest[i: i + spec.rounds_per_page])
    return [c for c in chunks if c]


def build_scorecard(doc, spec, *, placeholder_rounds=frozenset()):
    """Append one player's scorecard to ``doc`` (assumes the page is fresh).

    Returns the round tables. ``placeholder_rounds`` are the rounds to leave a
    seat placeholder in; the caller resolves those (per player) afterwards.
    """
    title = _add_paragraph(doc, spec.tournament_name, 25, bold=True)
    if spec.qr_url:
        _add_floating_image(title, BytesIO(_make_qr_png(spec.qr_url)),
                            QR_SIZE, QR_SIZE, QR_H_OFFSET, QR_V_OFFSET, behind=True)
    _add_floating_image(title, LOGO_PATH, LOGO_SIZE, LOGO_SIZE,
                        LOGO_H_OFFSET, LOGO_V_OFFSET, behind=False)

    _add_paragraph(doc, spec.tournament_date, 20)
    for _ in range(3):
        _add_paragraph(doc, "", 11)
    _add_paragraph(doc, spec.player_name, 30, bold=True, center=False)
    _add_paragraph(doc, "", 11, center=False)

    chunks = _chunk_rounds(spec)
    tables = []
    for i, chunk in enumerate(chunks):
        tables.append(_add_round_table(doc, chunk, spec.opponents, spec.results, placeholder_rounds))
        if i < len(chunks) - 1:
            _add_page_break(doc)

    _add_paragraph(doc, "", 11)
    _add_paragraph(doc, "", 11)
    footer = _add_paragraph(doc, "", 14)
    _set_run_font(footer.add_run(spec.footer_text + "\n"), 14)
    _set_run_font(footer.add_run(spec.footer_url), 14, bold=True)
    _add_paragraph(doc, "", 11)
    return tables


def _configure_page(section):
    section.page_width = Emu(7772400)    # 8.5"
    section.page_height = Emu(10058400)  # 11"
    section.left_margin = Emu(594360)
    section.right_margin = Emu(548640)
    section.top_margin = Emu(548640)
    section.bottom_margin = Emu(274320)


# Placeholder dropped into the template's player-name run, then swapped per clone.
_NAME_SENTINEL = "PLAYER_NAME"


def _opp_sentinel(round_number):
    return f"OPPONENT_{round_number}"


def _layout_signature(spec):
    """What fixes a scorecard's structure: everything but the per-player values
    (player name, opponent names, first/second marks) patched into clones."""
    return replace(spec, player_name="", opponents={}, starts={})


def _patch_text(element, replacements, spec):
    """Apply every per-player edit in a single pass over ``element``'s text.

    Sentinel ``<w:t>`` strings (player name, opponents) are swapped in place;
    start placeholders are collected and, after the walk, each is turned into an
    ellipse over the right ordinal or removed (see :func:`_resolve_start`).
    Folding both into one traversal means locating the seat marks is free — it
    rides the walk the name/opponent patch already makes.
    """
    starts = []
    for t in element.iter(qn("w:t")):
        text = t.text
        if not text:
            continue
        if text in replacements:
            t.text = replacements[text]
        elif text.startswith(_START_PREFIX):
            starts.append((t.getparent(), int(text[len(_START_PREFIX):])))
    for run, round_number in starts:
        _resolve_start(run, round_number, spec)


def _reassign_drawing_ids(element):
    """Give cloned drawings fresh ids so Word doesn't see duplicate object ids.

    Both the wordprocessing docPr id (wp:docPr/@id) and the DrawingML picture id
    (pic:cNvPr/@id) must be unique across the package or Word for the web treats
    the file as corrupt; the clone starts out sharing the template's ids for
    both.
    """
    for docPr in element.iter(qn("wp:docPr")):
        _uid[0] += 1
        docPr.set("id", str(_uid[0]))
        docPr.set("name", f"image{_uid[0]}")
    for cnvpr in element.iter(qn("pic:cNvPr")):
        _uid[0] += 1
        cnvpr.set("id", str(_uid[0]))


def build_document(specs):
    """Build a Document with one scorecard per spec, each starting a new page.

    Within a division every scorecard shares the same structure, QR and logo;
    only the player name and any prefilled opponent names differ. So we build
    the first one as a template (with placeholders in those spots) and deep-copy
    its XML for the rest, patching only the placeholders. That skips almost all
    per-element construction and renders the QR / reads the logo once. Specs
    whose layout differs from the first fall back to a fresh build.
    """
    doc = Document()
    _configure_page(doc.sections[0])
    body = doc.element.body
    for p in list(doc.paragraphs):
        body.remove(p._element)
    sectPr = body.find(qn("w:sectPr"))

    if not specs:
        return doc

    # Rounds that any player has a prefilled opponent for get a text placeholder
    # in the template so every clone can patch its own name into that slot;
    # rounds that any player has a fixed seat for get a seat placeholder, which
    # each clone resolves into an ellipse (or removes). Both are handled in the
    # clone's single patch pass.
    prefill_rounds = {n for spec in specs for n in spec.opponents}
    template_opponents = {n: _opp_sentinel(n) for n in prefill_rounds}
    start_rounds = frozenset(n for spec in specs for n in spec.starts)
    build_scorecard(doc, replace(
        specs[0], player_name=_NAME_SENTINEL,
        opponents=template_opponents, starts={},
    ), placeholder_rounds=start_rounds)
    template_els = [c for c in body if c is not sectPr]
    for el in template_els:
        body.remove(el)

    def append(el):
        sectPr.addprevious(el) if sectPr is not None else body.append(el)

    template_sig = _layout_signature(specs[0])
    for i, spec in enumerate(specs):
        if i > 0:
            _add_page_break(doc)
        if _layout_signature(spec) == template_sig:
            replacements = {_NAME_SENTINEL: spec.player_name}
            for n in prefill_rounds:
                replacements[_opp_sentinel(n)] = spec.opponents.get(n, "")
            for el in template_els:
                clone = copy.deepcopy(el)
                _reassign_drawing_ids(clone)
                _patch_text(clone, replacements, spec)
                append(clone)
        else:
            # Different layout: build it fresh, with seat placeholders for its
            # own rounds, then resolve them (no name/opponent sentinels here).
            tables = build_scorecard(doc, spec, placeholder_rounds=frozenset(spec.starts))
            for table in tables:
                _patch_text(table._tbl, {}, spec)
    return doc


def render_scorecards(specs):
    """Render the scorecards for ``specs`` and return the .docx file as bytes."""
    buf = BytesIO()
    build_document(specs).save(buf)
    return buf.getvalue()
